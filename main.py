import os
from docx import Document
import random

# Path to your 'spanish_docs' folder
folder_path = 'spanish_docs'

def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def parse_table_from_word(file_path):
    """
    Extracts vocabulary from Word tables if present.
    Assumes first and last columns contain English and Spanish respectively.
    """
    doc = Document(file_path)
    vocab_dict = {}

    # Check for tables in the document
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                # Skip header row and extract English and Spanish words
                english_word = row.cells[0].text.strip()
                spanish_word = row.cells[-1].text.strip()
                if english_word and spanish_word:
                    vocab_dict[spanish_word] = english_word
    return vocab_dict

# Reading all Word files from the 'spanish_docs' folder
files = [
    'Trifold-L01-nouns(1).docx',
    'Trifold-L01-numbers(1).docx',
    'Trifold-L01-pronouns.ser.time(1)(1).docx',
    'Trifold-L02-AR-Verbs(1)(1).docx',
    'Trifold-L02-Interrogatives(1)(1).docx',
    'Trifold-L02-Nouns-Portales2.0(1)(1).docx'
]

# Read and store vocabulary from all files
all_vocab = {}
for file in files:
    file_path = os.path.join(folder_path, file)
    vocab = parse_table_from_word(file_path)
    if vocab:
        all_vocab[file] = vocab
    else:
        print(f"No vocabulary found in {file}. Skipping.")

def generate_random_quiz(all_vocab):
    file_choice = random.choice(list(all_vocab.keys()))
    vocab_dict = all_vocab[file_choice]

    # Check if there are enough words for options
    if len(vocab_dict) < 3:
        print(f"Not enough words in {file_choice} for quiz. Skipping.")
        return None

    # Select a random question and its correct answer
    question = random.choice(list(vocab_dict.keys()))
    correct_answer = vocab_dict[question]

    # Create a list of wrong answers and ensure they don't include the correct one
    wrong_answers = list(vocab_dict.values())
    wrong_answers.remove(correct_answer)  # Remove correct answer from wrong answers pool
    options = random.sample(wrong_answers, min(3, len(wrong_answers)))  # Select wrong options

    # Add the correct answer to options and shuffle them
    options.append(correct_answer)
    random.shuffle(options)

    # Display the quiz
    print(f"\nFrom {file_choice}: What is the meaning of '{question}'?")
    for i, option in enumerate(options):
        print(f"{i + 1}. {option}")

    # Return the correct answer and options for user selection
    return correct_answer, options

def quiz():
    score = 0
    num_questions = 10

    for i in range(num_questions):
        result = generate_random_quiz(all_vocab)

        if result:
            correct_answer, options = result
            try:
                # Get user input and validate answer
                user_choice = int(input("Choose the correct answer (1-4): "))
                if options[user_choice - 1] == correct_answer:
                    print("Correct!")
                    score += 1
                else:
                    print(f"Wrong! The correct answer was: {correct_answer}")
            except (ValueError, IndexError):
                print("Invalid input. Please choose a number between 1 and 4.")
        else:
            # If a file doesn't have enough vocab, skip to next question
            continue

    # Calculate final score
    percentage = (score / num_questions) * 100
    print(f"\nYou got {score} out of {num_questions} correct.")
    print(f"Your score: {percentage}%")

# Start the quiz
quiz()
